# coding=utf-8
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.shared import RGBColor

firma = u'Фирма ЕООД'
slujitel = u'Служител'
upravitel = u'Управител, ЕГН хххххххххх - длъжност'
schetovoditel = u'Счетоводител'

# Za komandirovkata
zapoved_nomer = 1
data = u'01.06.2019г.'
start_date = '02/10/2019'
end_date = '07/10/2019'
delta_days = 6
country = u'Държава'
city = u'Град'
vid_transport = u'Самолет'

cel = u'Среща с потенциален клиент'
uslovia = u'разходите по допълнителните разходи по транспорт на мястото на лицето са включени. Градският и ' \
          u'междуградският транспорт е включен.'
suma_na_den = 50
valuta = 'EUR'
kurs_1_valuta_kym_1_bgn = 1.9558

samolet_vav_valuta = 38.2
samolet_v_lv = round(kurs_1_valuta_kym_1_bgn*samolet_vav_valuta)

# hotel_vav_valuta = 150
# hotel_v_lv = round(hotel_vav_valuta*kurs_1_valuta_kym_1_bgn)

# taxi_vav_valuta = 49
# taxi_v_lv = round(taxi_vav_valuta * kurs_1_valuta_kym_1_bgn)

# telefon = 22
# telefon_v_lv = telefon*kurs_1_valuta_kym_1_bgn

# gradski_transport_vav_valuta = 20
# gradski_transport_v_lv = round(gradski_transport_vav_valuta*kurs_1_valuta_kym_1_bgn)

razhodi = (
    # ('Тип Разход', 'Валута','Курс','Сума','Разходи от работодател в лв.', 'Разходи от служител в лв.'),
    #        (u'Пътни - градски транспорт',valuta, kurs_1_valuta_kym_1_bgn, gradski_transport_vav_valuta, 0, gradski_transport_v_lv),
           (u'Пътни - самолетен билет',valuta, kurs_1_valuta_kym_1_bgn, samolet_vav_valuta, 0, samolet_v_lv),
    #        (u'Хотел',valuta, kurs_1_valuta_kym_1_bgn, hotel_vav_valuta, 0, hotel_v_lv),
    #        (u'Хотел - 3',valuta, kurs_1_valuta_kym_1_bgn, hotel_3, 0, hotel_3_v_lv),
    #        (u'Телефон - Сим карта', valuta, kurs_1_valuta_kym_1_bgn, telefon, 0, telefon_v_lv),
    (u'Дневни', valuta, kurs_1_valuta_kym_1_bgn, suma_na_den * delta_days, 0, round(suma_na_den * delta_days * kurs_1_valuta_kym_1_bgn))
)

# Variables
total_suma_vav_valuta = delta_days * suma_na_den
total_suma_v_bgn = total_suma_vav_valuta * kurs_1_valuta_kym_1_bgn

folder = ''


# Functions
def generate_zapoved():
    document = Document()
    run = Document().add_paragraph().add_run()
    font = run.font
    black_color = RGBColor(0x00, 0x00, 0x00)
    font.color.rgb = black_color

    heading = document.add_heading(firma, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    zapoved_data = u"Заповед No. {0} /   {1}".format(str(zapoved_nomer), data)
    paragraph_zapoved = document.add_heading(level=1)
    paragraph_zapoved.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_zapoved_run = paragraph_zapoved.add_run(zapoved_data)
    paragraph_zapoved_run.font.color.rgb = black_color

    p = document.add_paragraph(u'\nНа основание Наредба за служебните командировки и специализации в чужбина')
    p = document.add_paragraph('\n')
    p.add_run(upravitel).underline = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = document.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(u'Нареждам:')
    heading_run.font.color.rgb = black_color
    heading.bold = True

    p = document.add_paragraph(u'\nКомандировам ')
    p.add_run(upravitel).underline = True

    p = document.add_paragraph(u'За периода от ')
    p.add_run(start_date).underline = True
    p.add_run(u' до ')
    p.add_run(end_date).underline = True
    p.add_run(u' продължителност {0} дни в {1}, {2}'.format(delta_days, country, city))
    p.add_run(u'\n\nС цел: ')
    p.add_run(cel).underline = True
    p.add_run(u'\n\nПри следните условия: ')
    p.add_run(uslovia).underline = True
    p.add_run(
        u'\n\nДа се отпуснат дневни размер на {0} дни х {1} {2} = {3} {2} или {4} лв.'.format(delta_days, suma_na_den,
                                                                                              valuta,
                                                                                              total_suma_vav_valuta,
                                                                                              total_suma_v_bgn))
    p = document.add_paragraph('\n')
    p_run = p.add_run(u'\nВсичко: {0} лв.'.format(total_suma_v_bgn))
    p_run.bold = True
    p_run.underline = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = document.add_paragraph(u'\nКопие от заповедта да се връчи на лицата и главния счетоводител.')
    p = document.add_paragraph(u'............................................... - управител')
    p = document.add_paragraph(u'\n\nДата: {0}'.format(data))

    document.save(folder + '01_Zapoved_{0}.docx'.format(zapoved_nomer))


def generate_lichen_finansov_otchet():
    document = Document()
    run = Document().add_paragraph().add_run()
    font = run.font
    black_color = RGBColor(0x00, 0x00, 0x00)
    font.color.rgb = black_color

    heading = document.add_heading(firma, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    zapoved_data = u"Личен финансов отчет за командировка"
    paragraph_zapoved = document.add_heading(level=1)
    paragraph_zapoved.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_zapoved_run = paragraph_zapoved.add_run(zapoved_data)
    paragraph_zapoved_run.font.color.rgb = black_color

    p = document.add_paragraph(u'\nот командировка в чужбина')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(u'От: ')
    p.add_run(slujitel).underline = True
    document.add_paragraph(u'Отдел: Информационни технологии')
    document.add_paragraph(u'В /държава, град/: {0}, {1}'.format(country, city))
    document.add_paragraph(u'За период /от-до/: {0} - {1}'.format(start_date, end_date))
    document.add_paragraph(u'Цел на командировката: {0}'.format(cel))
    document.add_paragraph(u'Вид транспорт: {0}'.format(vid_transport))
    document.add_paragraph(u'Валута: {0}. Курс: 1 {0} = {1} лева'.format(valuta, kurs_1_valuta_kym_1_bgn))

    generate_table(document)

    p = document.add_paragraph()
    p.add_run(u'\n\n\nКомандирован /{0}/: ...............................................'.format(slujitel)).bold = True

    p.add_run(u'\n\nДата: {0}.{1}.{2}г.'.format(date.today().day, date.today().month, date.today().year))

    document.save(folder + '02_Lichen_Fianansov_otchet_{0}.docx'.format(zapoved_nomer))


def generate_finansov_otchet():
    document = Document()
    run = Document().add_paragraph().add_run()
    font = run.font
    black_color = RGBColor(0x00, 0x00, 0x00)
    font.color.rgb = black_color

    heading = document.add_heading(firma, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    zapoved_data = u"Финансов отчет"
    paragraph_zapoved = document.add_heading(level=1)
    paragraph_zapoved.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_zapoved_run = paragraph_zapoved.add_run(zapoved_data)
    paragraph_zapoved_run.font.color.rgb = black_color

    p = document.add_paragraph(u'\nот командировка в чужбина')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(u'От: ')
    p.add_run(upravitel).underline = True
    document.add_paragraph(u'Отдел: Информационни технологии')
    document.add_paragraph(u'В /държава, град/: {0}, {1}'.format(country, city))
    document.add_paragraph(u'За период /от-до/: {0} - {1}'.format(start_date, end_date))
    document.add_paragraph(u'Цел на командировката: {0}'.format(cel))

    (razhodi_rabotodatel, razhodi_slujitel) = generate_table(document)

    document.add_paragraph(u'Общо признати разходи за сметка на работодател: {0}лв.'.format(razhodi_rabotodatel))
    p = document.add_paragraph(u'Общо признати разходи за сметка на служител: {0}лв.'.format(razhodi_slujitel))
    p.add_run(u'\n\nСума за възстановяване на служител: {0}лв.'.format(razhodi_slujitel)).bold = True
    p.add_run(u'\n\n\nУправител /{0}/: ...............................................'.format(upravitel)).bold = True

    p.add_run(u'\n\nСъгласувано с:')
    p.add_run(u'\n\nСчетоводител /{0}/: ...............................................'.format(schetovoditel)).bold = True

    p.add_run(u'\n\nДата: {0}.{1}.{2}г.'.format(date.today().day, date.today().month, date.today().year))

    document.save(folder + '03_Fianansov_otchet_{0}.docx'.format(zapoved_nomer))


def generate_table(document):
    razhodi_slujitel = 0
    razhodi_rabotodatel = 0

    table = document.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'Тип Разход'
    hdr_cells[1].text = u'Валута'
    hdr_cells[2].text = u'Курс'
    hdr_cells[3].text = u'Сума'
    hdr_cells[4].text = u'Разходи, платени от работодател в лв.'
    hdr_cells[5].text = u'Разходи, платени от служител в лв.'

    for row in razhodi:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = str(row[2])
        row_cells[3].text = str(row[3])
        row_cells[4].text = str(row[4])
        row_cells[5].text = str(row[5]) + u' лв.'

        razhodi_rabotodatel += row[4]
        razhodi_slujitel += row[5]

    row_cells = table.add_row().cells
    row_cells[0].text = u'Общо'
    row_cells[4].text = str(razhodi_rabotodatel)
    row_cells[5].text = str(razhodi_slujitel) + u' лв.'

    return (razhodi_rabotodatel, razhodi_slujitel)


generate_zapoved()
generate_lichen_finansov_otchet()
generate_finansov_otchet()
